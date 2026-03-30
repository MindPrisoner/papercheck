import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.config import DATA_DIR

PARSED_DIR = DATA_DIR / "parsed"
PARSED_DIR.mkdir(parents=True, exist_ok=True)


def clean_line(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\u3000", " ")
    text = text.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_key(text: str) -> str:
    if not text:
        return ""
    text = text.strip().lower()
    text = text.replace("：", ":")
    text = re.sub(r"\s+", "", text)
    return text


def extract_docx_blocks(file_path: str) -> list[dict]:
    """
    提取 docx 中的“文本块”。
    先读普通段落，再补表格中的文本。
    对毕业论文 MVP 来说，这样已经够用了。
    """
    doc = DocxDocument(file_path)
    blocks: list[dict] = []

    for p in doc.paragraphs:
        text = clean_line(p.text)
        style_name = p.style.name if p.style else ""
        if text:
            blocks.append(
                {
                    "text": text,
                    "style": style_name,
                }
            )

    # 补充表格里的内容
    # 注意：这不一定严格保持原始顺序，但能提升命中率
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_parts = []
                for p in cell.paragraphs:
                    t = clean_line(p.text)
                    if t:
                        cell_parts.append(t)
                if cell_parts:
                    row_parts.append(" ".join(cell_parts))
            if row_parts:
                blocks.append(
                    {
                        "text": " | ".join(row_parts),
                        "style": "Table",
                    }
                )

    return blocks


def extract_pdf_blocks(file_path: str) -> list[dict]:
    """
    PDF 版本先做最小实现。
    PyMuPDF 文档说明 sort=True 会更接近自然阅读顺序。
    """
    pdf = fitz.open(file_path)
    blocks: list[dict] = []

    for page in pdf:
        text = page.get_text("text", sort=True)
        for line in text.splitlines():
            t = clean_line(line)
            if t:
                blocks.append(
                    {
                        "text": t,
                        "style": "PDF",
                    }
                )

    return blocks


def load_blocks(file_path: str) -> list[dict]:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".docx":
        return extract_docx_blocks(file_path)
    elif suffix == ".pdf":
        return extract_pdf_blocks(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {suffix}")


def blocks_to_raw_text(blocks: list[dict]) -> str:
    return normalize_text("\n".join(block["text"] for block in blocks if block["text"]))


def save_raw_text(public_id: str, text: str) -> str:
    raw_text_path = PARSED_DIR / f"{public_id}.txt"
    raw_text_path.write_text(text, encoding="utf-8")
    return str(raw_text_path)


def is_abstract_heading(text: str) -> bool:
    key = normalize_key(text)
    return key in {"摘要", "abstract", "中文摘要", "英文摘要"}


def is_keywords_heading(text: str) -> bool:
    key = normalize_key(text)
    return key in {"关键词", "关键字", "keywords"}


def is_reference_heading(text: str) -> bool:
    key = normalize_key(text)
    return key in {"参考文献", "references", "reference"}


def extract_title(blocks: list[dict]) -> str | None:
    """
    简单标题策略：
    在前 12 个文本块里找第一个像标题的内容。
    """
    if not blocks:
        return None

    candidates = [b["text"] for b in blocks[:12] if b["text"]]

    for line in candidates:
        key = normalize_key(line)
        if key in {"摘要", "abstract", "关键词", "关键字", "目录"}:
            continue
        if 2 <= len(line) <= 80:
            return line

    return candidates[0] if candidates else None


def extract_inline_labeled_text(text: str, patterns: list[str]) -> str | None:
    for pattern in patterns:
        m = re.match(pattern, text, re.I)
        if m:
            value = m.group(1).strip()
            if value:
                return value
    return None


def extract_abstract(blocks: list[dict]) -> str | None:
    """
    支持两种常见格式：
    1. 摘要：xxxx
    2. 单独一行“摘 要/摘要”，后面几段是摘要正文
    """
    inline_patterns = [
        r"^\s*摘\s*要\s*[:：]\s*(.+)$",
        r"^\s*abstract\s*[:：]\s*(.+)$",
    ]

    # 先找“摘要：正文”这种行内格式
    for block in blocks[:60]:
        text = block["text"]
        inline = extract_inline_labeled_text(text, inline_patterns)
        if inline:
            return inline[:3000]

    # 再找单独摘要标题
    for i, block in enumerate(blocks[:120]):
        text = block["text"]

        if is_abstract_heading(text):
            abstract_parts = []

            for nxt in blocks[i + 1:]:
                nxt_text = nxt["text"]

                if is_keywords_heading(nxt_text):
                    break
                if is_reference_heading(nxt_text):
                    break
                if is_section_heading(nxt_text, nxt.get("style", "")):
                    break

                abstract_parts.append(nxt_text)

            abstract_text = "\n".join(abstract_parts).strip()
            if abstract_text:
                return abstract_text[:3000]

    return None


def split_keywords(raw: str) -> list[str]:
    if not raw:
        return []

    parts = re.split(r"[，,；;、]+", raw)
    result = []

    for part in parts:
        t = clean_line(part)
        if t:
            result.append(t)

    return result[:8]


def extract_keywords(blocks: list[dict]) -> list[str]:
    inline_patterns = [
        r"^\s*关\s*键\s*词\s*[:：]\s*(.+)$",
        r"^\s*关\s*键\s*字\s*[:：]\s*(.+)$",
        r"^\s*keywords\s*[:：]\s*(.+)$",
    ]

    # 先找“关键词：xxx”
    for block in blocks[:120]:
        text = block["text"]
        inline = extract_inline_labeled_text(text, inline_patterns)
        if inline:
            return split_keywords(inline)

    # 再找单独一行“关键词”，下一行是内容
    for i, block in enumerate(blocks[:120]):
        text = block["text"]
        if is_keywords_heading(text):
            if i + 1 < len(blocks):
                return split_keywords(blocks[i + 1]["text"])

    return []


def is_section_heading(text: str, style: str = "") -> bool:
    text = clean_line(text)
    if not text:
        return False

    if len(text) > 60:
        return False

    if is_abstract_heading(text) or is_keywords_heading(text) or is_reference_heading(text):
        return True

    style_lower = (style or "").lower()
    if "heading" in style_lower:
        return True

    patterns = [
        r"^第[一二三四五六七八九十百零\d]+章\s*.+$",
        r"^[一二三四五六七八九十]+、.+$",
        r"^\d+\s+\S.+$",
        r"^\d+\.\d+\s*\S.*$",
        r"^\d+\.\d+\.\d+\s*\S.*$",
        r"^\([一二三四五六七八九十\d]+\)\s*.+$",
        r"^（[一二三四五六七八九十\d]+）\s*.+$",
    ]

    for pattern in patterns:
        if re.match(pattern, text):
            return True

    return False


def infer_heading_level(text: str, style: str = "") -> int:
    style_lower = (style or "").lower()
    if "heading 1" in style_lower:
        return 1
    if "heading 2" in style_lower:
        return 2
    if "heading 3" in style_lower:
        return 3

    text = clean_line(text)

    if re.match(r"^第[一二三四五六七八九十百零\d]+章", text):
        return 1
    if re.match(r"^[一二三四五六七八九十]+、", text):
        return 1

    m = re.match(r"^(\d+(?:\.\d+)*)", text)
    if m:
        return m.group(1).count(".") + 1

    if re.match(r"^\([一二三四五六七八九十\d]+\)", text) or re.match(r"^（[一二三四五六七八九十\d]+）", text):
        return 2

    return 1


def split_sections(blocks: list[dict]) -> list[dict]:
    """
    更适合中文论文的章节切分：
    - 支持 Heading 样式
    - 支持 1 / 1.1 / 1.1.1
    - 支持 第一章 / 一、
    """
    sections = []
    current_heading = None
    current_level = 1
    current_content: list[str] = []
    started = False

    def flush_section():
        nonlocal current_heading, current_level, current_content, sections
        if current_heading and current_content:
            content = "\n".join(current_content).strip()
            if content:
                sections.append(
                    {
                        "heading": current_heading,
                        "level": current_level,
                        "content": content[:8000],
                    }
                )

    for block in blocks:
        text = block["text"]
        style = block.get("style", "")

        if is_reference_heading(text):
            break

        # 摘要和关键词不算正文章节
        if is_abstract_heading(text) or is_keywords_heading(text):
            continue

        if is_section_heading(text, style):
            # 跳过“摘要/关键词/参考文献”这些非正文标题
            if is_abstract_heading(text) or is_keywords_heading(text) or is_reference_heading(text):
                continue

            if started:
                flush_section()

            current_heading = text
            current_level = infer_heading_level(text, style)
            current_content = []
            started = True
        else:
            if started:
                current_content.append(text)

    if started:
        flush_section()

    # 如果一篇文档完全没识别到标题，就给一个兜底正文
    if not sections:
        body_lines = []
        for block in blocks:
            text = block["text"]
            if is_abstract_heading(text) or is_keywords_heading(text) or is_reference_heading(text):
                continue
            body_lines.append(text)

        body_text = "\n".join(body_lines).strip()
        if body_text:
            sections.append(
                {
                    "heading": "正文",
                    "level": 1,
                    "content": body_text[:8000],
                }
            )

    return sections[:40]


def is_new_reference_line(text: str) -> bool:
    patterns = [
        r"^\[\d+\]",
        r"^\d+\.",
        r"^\d+\s",
        r"^[（(]\d+[）)]",
    ]
    for pattern in patterns:
        if re.match(pattern, text):
            return True
    return False


def extract_references(blocks: list[dict]) -> list[str]:
    start_idx = None

    for i, block in enumerate(blocks):
        if is_reference_heading(block["text"]):
            start_idx = i + 1
            break

    if start_idx is None:
        return []

    ref_lines = []
    for block in blocks[start_idx:]:
        text = clean_line(block["text"])
        if text:
            ref_lines.append(text)

    if not ref_lines:
        return []

    refs: list[str] = []
    for line in ref_lines:
        if not refs:
            refs.append(line)
            continue

        if is_new_reference_line(line):
            refs.append(line)
        else:
            refs[-1] += " " + line

    cleaned = []
    for ref in refs:
        ref = clean_line(ref)
        if len(ref) >= 6:
            cleaned.append(ref)

    return cleaned[:100]


def parse_document(file_path: str, public_id: str) -> dict:
    blocks = load_blocks(file_path)
    raw_text = blocks_to_raw_text(blocks)

    title = extract_title(blocks)
    abstract_text = extract_abstract(blocks)
    keywords = extract_keywords(blocks)
    sections = split_sections(blocks)
    references = extract_references(blocks)
    raw_text_path = save_raw_text(public_id, raw_text)

    word_count = len(re.sub(r"\s+", "", raw_text))
    section_count = len(sections)
    reference_count = len(references)

    return {
        "title": title,
        "abstract_text": abstract_text,
        "keywords": keywords,
        "sections": sections,
        "references": references,
        "raw_text_path": raw_text_path,
        "word_count": word_count,
        "section_count": section_count,
        "reference_count": reference_count,
        "raw_text_preview": raw_text[:1500],
    }